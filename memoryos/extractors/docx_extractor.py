from docx import Document

from memoryos.extractors.result import ExtractionResult


def extract_docx(path) -> ExtractionResult:
    document = Document(path)
    text_parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return ExtractionResult(
        text="\n".join(text_parts),
        structural_metadata={"paragraph_count": len(document.paragraphs)},
    )
