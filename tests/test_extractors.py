from pathlib import Path

import pytest

from memoryos.extractors.docx_extractor import extract_docx
from memoryos.extractors.pdf_extractor import extract_pdf
from memoryos.extractors.pptx_extractor import extract_pptx
from memoryos.extractors.xlsx_extractor import extract_xlsx
from memoryos.ocr.tesseract_engine import TesseractOcrEngine

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEST_FILES_2 = PROJECT_ROOT / "all test data" / "test files 2"
TEST_FILES_3 = PROJECT_ROOT / "all test data" / "test files 3"


@pytest.fixture(scope="module")
def ocr_engine():
    return TesseractOcrEngine()


def test_extract_pdf_gets_text_and_page_count(ocr_engine):
    result = extract_pdf(TEST_FILES_2 / "world-map.pdf", ocr_engine)
    assert result.structural_metadata["page_count"] == 1
    assert "Greenland" in result.text or "Brazil" in result.text


def test_extract_pdf_caps_embedded_images(ocr_engine):
    result = extract_pdf(TEST_FILES_2 / "first aid.pdf", ocr_engine)
    assert result.structural_metadata["page_count"] == 4
    assert len(result.embedded_images) <= 5
    assert result.text.strip() != ""


def test_extract_pptx_gets_text_and_slide_count():
    result = extract_pptx(TEST_FILES_3 / "70YEARS ISRAEL.PPTX")
    assert result.structural_metadata["slide_count"] > 0
    assert "ישראל" in result.text
    assert len(result.embedded_images) <= 5


def test_extract_docx_roundtrip(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Quarterly financial report for the coffee division")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Region"
    table.rows[0].cells[1].text = "Revenue"
    docx_path = tmp_path / "smoketest.docx"
    doc.save(docx_path)

    result = extract_docx(docx_path)
    assert "Quarterly financial report" in result.text
    assert "Region" in result.text
    assert "Revenue" in result.text
    assert result.structural_metadata["paragraph_count"] == 1


def test_extract_xlsx_roundtrip(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws["A1"] = "Product"
    ws["B1"] = "Green coffee beans"
    xlsx_path = tmp_path / "smoketest.xlsx"
    wb.save(xlsx_path)

    result = extract_xlsx(xlsx_path)
    assert "Green coffee beans" in result.text
    assert result.structural_metadata["sheet_names"] == ["Sales"]
